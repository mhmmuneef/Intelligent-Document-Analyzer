from flask import Flask, request, redirect, url_for, render_template
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
import docx
import openai
from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.orm import declarative_base, sessionmaker
from dotenv import load_dotenv
from contextlib import contextmanager
from PyPDF2.errors import PdfReadError
import json
import os
import urllib

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'

# Ensure uploads folder exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Load environment variables
load_dotenv()

# Azure OpenAI Configuration
openai.api_type = "azure"
openai.api_base = os.getenv('AZURE_OPENAI_ENDPOINT').rstrip('/')
openai.api_version = "2023-05-15"  # Use the appropriate API version
openai.api_key = os.getenv('AZURE_OPENAI_KEY')
deployment_name = os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME')

# Database Connection String
DATABASE_URI = os.getenv('DATABASE_URI')

# Initialize Database
params = urllib.parse.quote_plus(DATABASE_URI)
engine = create_engine(f"mssql+pyodbc:///?odbc_connect={params}")
Session = sessionmaker(bind=engine)
Base = declarative_base()

# Define the Document model
class Document(Base):
    __tablename__ = 'documents'
    id = Column(Integer, primary_key=True)
    filename = Column(String)
    content = Column(Text)
    entities = Column(Text)

# Create tables
Base.metadata.create_all(engine)

@contextmanager
def get_session():
    session = Session()
    try:
        yield session
    finally:
        session.close()

def extract_text_from_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        return text
    except PdfReadError as e:
        print(f"PDF read error: {e}")
        return ''
    except Exception as e:
        print(f"An error occurred while extracting text from PDF: {e}")
        return ''

def extract_text_from_docx(filepath):
    doc = docx.Document(filepath)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

def extract_text(filepath):
    if filepath.lower().endswith('.pdf'):
        return extract_text_from_pdf(filepath)
    elif filepath.lower().endswith('.docx'):
        return extract_text_from_docx(filepath)
    else:
        return ''

def analyze_text(text):
    prompt = (
        "You are an AI assistant that extracts key entities from text. "
        "Please extract the names of organizations, people, locations, dates, and contact information from the following text:\n"
        f"{text}\n\n"
        "Return the results in JSON format with the fields: 'text', 'category', and 'confidence_score'."
    )

    try:
        response = openai.ChatCompletion.create(
            engine=deployment_name,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500,
            temperature=0,
        )
        # Get the content from the response
        response_content = response.choices[0].message['content']
        print("Raw API Response:", response_content)  # Debugging line
        
        # Clean up the response by removing the code block formatting
        json_str = response_content.strip('```json\n').strip('```')

        # Parse the JSON
        entities = json.loads(json_str)
        return entities
    except json.JSONDecodeError as e:
        print(f"JSON decoding error: {e}. Response content was: {response_content}")
        return []
    except Exception as e:
        print(f"OpenAI API error: {e}")
        return []


ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'document' not in request.files:
        return 'No file part', 400
    file = request.files['document']
    if file.filename == '':
        return 'No selected file', 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract text and analyze
        text = extract_text(filepath)
        print("Extracted Text:", text)  # Debugging line

        entities = analyze_text(text)
        print("Extracted Entities:", entities)  # Debugging line

        # Save to database
        with get_session() as session:
            try:
                document = Document(
                    filename=filename,
                    content=text,
                    entities=json.dumps(entities)
                )
                session.add(document)
                session.commit()
                return redirect(url_for('result', document_id=document.id))
            except Exception as e:
                session.rollback()
                print(f"Database error: {e}")
                return 'Internal Server Error', 500
    else:
        return 'Unsupported file type', 400

@app.route('/result/<int:document_id>')
def result(document_id):
    with get_session() as session:
        document = session.query(Document).filter_by(id=document_id).first()
        if document:
            entities = json.loads(document.entities)
            return render_template('result.html', entities=entities)
        else:
            return 'Document not found', 404

if __name__ == '__main__':
    app.run(debug=True)
